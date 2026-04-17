import os
import logging
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from fpdf import FPDF
import unidecode

from security_gate.core.config import settings
from security_gate.core.risk_engine import RiskEngine

logger = logging.getLogger(__name__)

class ReportExporter:
    """
    Module Giai đoạn 4: Xuất báo cáo trực tiếp ra PDF, Word, Excel.
    Đã tích hợp: 
    - Nổi bật "Điểm mù dữ liệu" (Blind Spots).
    - Khôi phục chính xác Text báo cáo YARA & Hành vi.
    """
    
    def __init__(self, risk_engine: RiskEngine):
        self.engine = risk_engine
        self.report_dir = settings.REPORT_DIR
        self.report_dir.mkdir(parents=True, exist_ok=True)

    def _extract_report_data(self, analysis_data: Dict[str, Any]) -> Dict[str, Any]:
        r_total = analysis_data.get('final_score', 0.0)
        details = analysis_data.get('details', [])
        
        # 1. Điểm mù dữ liệu
        blind_spots = [comp for comp in details if comp.get('is_blind_spot', False)]
        has_blind_spot = len(blind_spots) > 0
        
        # 2. Phán quyết Human-in-the-loop
        if r_total >= 8.0:
            status = "🟠 CHỜ DUYỆT THỦ CÔNG (PENDING_REVIEW)" if has_blind_spot else "🟢 TIN DÙNG (APPROVED)"
        elif r_total >= 5.0:
            status = "🟡 CẢNH BÁO (PENDING)"
        else:
            status = "🔴 TỪ CHỐI (REJECTED)"

        w_v, w_m, w_i, w_l = self.engine.weights
        cr_score = self.engine.cr_score
        weakest_link_data = analysis_data.get('weakest_link', {})
        weakest = weakest_link_data.get('name', 'N/A') if isinstance(weakest_link_data, dict) else str(weakest_link_data)

        # Đánh dấu tên hiển thị
        for comp in details:
            if comp.get('is_blind_spot', False):
                comp['name_display'] = f"{comp.get('name', 'Unknown')} (⚠️ Mù dữ liệu)"
            else:
                comp['name_display'] = comp.get('name', 'Unknown')

        # 3. KHÔI PHỤC VĂN BẢN CẢNH BÁO MÃ ĐỘC
        alerts = []
        if has_blind_spot:
            alerts.append(f"⚠️ LƯU Ý: Hệ thống rớt mạng. Có {len(blind_spots)} thư viện cần tra cứu lỗ hổng thủ công.")

        integrity = analysis_data.get('project_integrity', {})
        
        if integrity.get('yara_alerts', 0) > 0:
            alerts.append(f"Cảnh báo YARA: Đã match {integrity['yara_alerts']} rule mã độc.")
            
        if integrity.get('malicious_vt', 0) > 0:
            alerts.append(f"Cảnh báo VirusTotal: Phát hiện {integrity['malicious_vt']} engine báo cáo mã độc.")
            
        for comp in details:
            if comp.get('heuristic_flag'):
                alerts.append(f"Cảnh báo Heuristic: Phát hiện logic rủi ro trong {comp.get('name')}.")
                
        if not alerts:
            alerts.append("Không phát hiện hành vi bất thường hoặc dấu vết mã độc.")

        return {
            "artifact": analysis_data.get('artifact', 'Unknown'),
            "hash": analysis_data.get('hash', 'Unknown'),
            "time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "status": status,
            "r_total": r_total,
            "weakest": weakest,
            "cr_score": cr_score,
            "weights": f"w_V={w_v:.2f}, w_M={w_m:.2f}, w_I={w_i:.2f}, w_L={w_l:.2f}",
            "details": details,
            "blind_spots": blind_spots,
            "alerts": alerts
        }

    def _generate_word(self, file_path: Path, data: dict):
        doc = Document()
        heading = doc.add_heading('BÁO CÁO KIỂM TOÁN CHUỖI CUNG ỨNG (ASVS 5.0.0)', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Artifact: {data['artifact']}")
        doc.add_paragraph(f"Hash (SHA256): {data['hash']}")
        doc.add_paragraph(f"Thời gian quét: {data['time']}")
        
        doc.add_heading('1. Kết quả Ra quyết định Tổng thể', level=2)
        doc.add_paragraph(f"- Trạng thái thực thi: {data['status']}")
        doc.add_paragraph(f"- Chỉ số Rủi ro tổng hợp (R): {data['r_total']} / 10.0")
        doc.add_paragraph(f"- Thành phần rủi ro nhất (Weakest Link): {data['weakest']}")
        
        if data['blind_spots']:
            doc.add_heading('🚨 DANH SÁCH THƯ VIỆN CẦN KIỂM TRA THỦ CÔNG (ĐIỂM MÙ MẠNG)', level=2)
            doc.add_paragraph("Hệ thống mất kết nối khi quét các thư viện sau. Vui lòng tra cứu lỗ hổng (CVE) thủ công dựa trên mã PURL:")
            for bs in data['blind_spots']:
                p = doc.add_paragraph(style='List Bullet')
                run_name = p.add_run(f"{bs.get('name')} (v{bs.get('version')})")
                run_name.bold = True
                run_name.font.color.rgb = RGBColor(220, 53, 69)
                p.add_run(f" | PURL: {bs.get('purl')}")

        doc.add_heading('2. Cơ sở Toán học (AHP & SAW)', level=2)
        cr_status = "✅ Đạt" if data['cr_score'] < 0.1 else "⚠️ Cần xem xét"
        doc.add_paragraph(f"- Tỷ số nhất quán (CR): {data['cr_score']:.4f} ({cr_status})")
        doc.add_paragraph(f"- Trọng số áp dụng (w_i): {data['weights']}")
        
        doc.add_heading('3. Chi tiết Điểm thành phần (Component Breakdown)', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'C_V'
        hdr_cells[2].text = 'C_M'
        hdr_cells[3].text = 'C_I'
        hdr_cells[4].text = 'C_L'
        hdr_cells[5].text = 'R_Score'
        
        for comp in data['details']:
            row_cells = table.add_row().cells
            c = comp.get('scores', {})
            version_str = f" (v{comp.get('version')})" if comp.get('version') else ""
            row_cells[0].text = f"{comp.get('name_display')}{version_str}"
            row_cells[1].text = str(c.get('CV', 0.0))
            row_cells[2].text = str(c.get('CM', 0.0))
            row_cells[3].text = str(c.get('CI', 0.0))
            row_cells[4].text = str(c.get('CL', 0.0))
            row_cells[5].text = str(comp.get('r_score', 0.0))
            
        doc.add_heading('4. Báo cáo Phân tích Hành vi & Mã độc (Heuristic/Malware)', level=2)
        for alert in data['alerts']:
            doc.add_paragraph(alert)
            
        doc.save(file_path)

    def _generate_excel(self, file_path: Path, data: dict):
        rows = []
        for comp in data['details']:
            c = comp.get('scores', {})
            version_str = f" (v{comp.get('version')})" if comp.get('version') else ""
            note = "⚠️ Cần kiểm tra tay (Mù mạng)" if comp.get('is_blind_spot') else "✅ Đã phân tích"
            
            rows.append({
                "Tên Thư viện": f"{comp.get('name')}{version_str}",
                "PURL (Định danh)": comp.get('purl', ''),
                "Lỗ hổng (C_V)": c.get('CV', 0.0),
                "Bảo trì (C_M)": c.get('CM', 0.0),
                "Toàn vẹn (C_I)": c.get('CI', 0.0),
                "Pháp lý (C_L)": c.get('CL', 0.0),
                "R_Score": comp.get('r_score', 0.0),
                "Trạng thái Dữ liệu": note
            })
        
        df = pd.DataFrame(rows) if rows else pd.DataFrame([{"Tên Thư viện": "N/A"}])
        
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Component Breakdown', index=False)
            worksheet = writer.sheets['Component Breakdown']
            worksheet.column_dimensions['B'].width = 40
            worksheet.column_dimensions['H'].width = 30
            
            summary_data = {
                "Thuộc tính": ["Artifact", "Hash", "Thời gian", "Trạng thái", "Điểm R", "Weakest Link", "Tỷ số CR", "Trọng số"],
                "Giá trị": [
                    data['artifact'], 
                    data['hash'], 
                    data['time'],
                    data['status'], 
                    data['r_total'],
                    data['weakest'],
                    data['cr_score'],
                    data['weights']
                ]
            }
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Summary', index=False)

    def _generate_pdf(self, file_path: Path, data: dict):
        pdf = FPDF()
        pdf.add_page()
        
        def st(txt): return unidecode.unidecode(str(txt))
            
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, st('BAO CAO KIEM TOAN CHUOI CUNG UNG (ASVS 5.0.0)'), ln=True, align='C')
        pdf.ln(5)
        
        pdf.set_font("Arial", '', 10)
        pdf.cell(0, 6, st(f"Artifact: {data['artifact']}"), ln=True)
        pdf.cell(0, 6, st(f"Hash: {data['hash']}"), ln=True)
        pdf.cell(0, 6, st(f"Thoi gian: {data['time']}"), ln=True)
        pdf.line(10, pdf.get_y()+2, 200, pdf.get_y()+2)
        pdf.ln(5)
        
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 8, st('1. Ket qua Tong the'), ln=True)
        pdf.set_font("Arial", '', 10)
        pdf.cell(0, 6, st(f"- Trang thai: {data['status']}"), ln=True)
        pdf.cell(0, 6, st(f"- Chi so Rui ro (R): {data['r_total']}"), ln=True)
        pdf.cell(0, 6, st(f"- Thanh phan rui ro nhat (Weakest Link): {data['weakest']}"), ln=True)
        pdf.ln(3)

        if data['blind_spots']:
            pdf.set_font("Arial", 'B', 11)
            pdf.set_text_color(220, 53, 69)
            pdf.cell(0, 8, st('>> YEU CAU KIEM TRA THU CONG (DIEM MU MANG)'), ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            pdf.multi_cell(0, 6, st("He thong mat ket noi mang. Vui long tra cuu CVE thu cong cho cac PURL sau:"))
            for bs in data['blind_spots']:
                pdf.cell(5)
                pdf.cell(0, 6, st(f"- {bs.get('name')} | PURL: {bs.get('purl')}"), ln=True)
            pdf.ln(3)
        
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 8, st('2. Co so Toan hoc (AHP & SAW)'), ln=True)
        pdf.set_font("Arial", '', 10)
        cr_status = "Dat" if data['cr_score'] < 0.1 else "Can xem xet"
        pdf.cell(0, 6, st(f"- Ty so nhat quan (CR): {data['cr_score']:.4f} ({cr_status})"), ln=True)
        pdf.cell(0, 6, st(f"- Trong so ap dung (w_i): {data['weights']}"), ln=True)
        pdf.ln(3)
        
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 8, st('3. Chi tiet Diem thanh phan'), ln=True)
        
        pdf.set_font("Arial", 'B', 9)
        col_widths = [90, 18, 18, 18, 18, 25] 
        headers = ['Component', 'C_V', 'C_M', 'C_I', 'C_L', 'R_Score']
        for i, h in enumerate(headers):
            pdf.cell(col_widths[i], 8, st(h), border=1, align='C')
        pdf.ln()
        
        pdf.set_font("Arial", '', 9)
        for comp in data['details']:
            c = comp.get('scores', {})
            name_ver = f"{comp.get('name_display')}"
            if len(name_ver) > 48: name_ver = name_ver[:45] + "..." 
            
            pdf.cell(col_widths[0], 8, st(name_ver), border=1)
            pdf.cell(col_widths[1], 8, str(c.get('CV', 0.0)), border=1, align='C')
            pdf.cell(col_widths[2], 8, str(c.get('CM', 0.0)), border=1, align='C')
            pdf.cell(col_widths[3], 8, str(c.get('CI', 0.0)), border=1, align='C')
            pdf.cell(col_widths[4], 8, str(c.get('CL', 0.0)), border=1, align='C')
            pdf.set_font("Arial", 'B', 9)
            pdf.cell(col_widths[5], 8, str(comp.get('r_score', 0.0)), border=1, align='C')
            pdf.set_font("Arial", '', 9)
            pdf.ln()
            
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 8, st('4. Bao cao Phan tich Hanh vi & Ma doc (Heuristic/Malware)'), ln=True)
        pdf.set_font("Arial", '', 10)
        for alert in data['alerts']:
            pdf.multi_cell(0, 6, st(alert)) 

        pdf.output(str(file_path))

    def export_report(self, analysis_data: Dict[str, Any], format_type: str) -> Path:
        hash_val = analysis_data.get('hash', 'unknown')[:10]
        filename = f"Audit_{hash_val}_{datetime.now().strftime('%H%M%S')}"
        data = self._extract_report_data(analysis_data)

        if format_type == "pdf":
            file_path = self.report_dir / f"{filename}.pdf"
            self._generate_pdf(file_path, data)
        elif format_type == "word":
            file_path = self.report_dir / f"{filename}.docx"
            self._generate_word(file_path, data)
        elif format_type == "excel":
            file_path = self.report_dir / f"{filename}.xlsx"
            self._generate_excel(file_path, data)
        else:
            raise ValueError("Định dạng không được hỗ trợ!")

        logger.info(f"[GĐ 4] Đã xuất báo cáo thành công tại: {file_path}")
        return file_path